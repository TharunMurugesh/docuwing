"""Native parser conformance tests for PDF, DOCX, XLSX, and CSV."""

from __future__ import annotations

import csv
import io

import docx
import fitz
import openpyxl
import pytest

from docuwing_engine.domain.entities import Document, SourceFormat
from docuwing_engine.parsers.csv.plugin import CsvParser
from docuwing_engine.parsers.docx.plugin import DocxParser
from docuwing_engine.parsers.pdf_text.plugin import PdfTextParser
from docuwing_engine.parsers.xlsx.plugin import XlsxParser
from docuwing_engine.testing.conformance.parser import ParserConformanceTests


class TestPdfTextParserConformance(ParserConformanceTests):
    def get_test_document(self) -> tuple[Document, io.BytesIO]:
        doc = Document(
            workspace="ws-test", filename="test.pdf", source_format=SourceFormat.PDF_TEXT
        )

        pdf = fitz.open()
        page = pdf.new_page()
        page.insert_text((50, 50), "Hello Docuwing PDF Parser")

        stream = io.BytesIO()
        pdf.save(stream)
        stream.seek(0)
        pdf.close()
        return doc, stream

    def create_plugin(self):
        return PdfTextParser()


class TestDocxParserConformance(ParserConformanceTests):
    def get_test_document(self) -> tuple[Document, io.BytesIO]:
        doc = Document(workspace="ws-test", filename="test.docx", source_format=SourceFormat.DOCX)

        docx_doc = docx.Document()
        docx_doc.add_heading("Title Heading", level=1)
        docx_doc.add_paragraph("This is a sample sentence in a DOCX paragraph.")

        stream = io.BytesIO()
        docx_doc.save(stream)
        stream.seek(0)
        return doc, stream

    def create_plugin(self):
        return DocxParser()


class TestXlsxParserConformance(ParserConformanceTests):
    def get_test_document(self) -> tuple[Document, io.BytesIO]:
        doc = Document(workspace="ws-test", filename="test.xlsx", source_format=SourceFormat.XLSX)

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(["Name", "Value"])
        ws.append(["Item A", 100])

        stream = io.BytesIO()
        wb.save(stream)
        stream.seek(0)
        return doc, stream

    def create_plugin(self):
        return XlsxParser()


class TestCsvParserConformance(ParserConformanceTests):
    def get_test_document(self) -> tuple[Document, io.BytesIO]:
        doc = Document(workspace="ws-test", filename="test.csv", source_format=SourceFormat.CSV)

        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["ID", "Name", "Score"])
        writer.writerow(["1", "Alpha", "95"])

        stream = io.BytesIO(output.getvalue().encode("utf-8"))
        return doc, stream

    def create_plugin(self):
        return CsvParser()


@pytest.mark.asyncio
async def test_native_parsers_structural_output():
    # Test PDF parser content
    pdf_test = TestPdfTextParserConformance()
    doc_pdf, stream_pdf = pdf_test.get_test_document()
    parser_pdf = PdfTextParser()
    ir_pdf = await parser_pdf.parse(doc_pdf, stream_pdf)
    assert len(ir_pdf.sections[0].blocks) >= 1
    assert "Hello Docuwing" in ir_pdf.sections[0].blocks[0].text

    # Test DOCX parser content
    docx_test = TestDocxParserConformance()
    doc_docx, stream_docx = docx_test.get_test_document()
    parser_docx = DocxParser()
    ir_docx = await parser_docx.parse(doc_docx, stream_docx)
    assert len(ir_docx.sections[0].blocks) == 2
    assert ir_docx.sections[0].blocks[0].role == "heading"

    # Test XLSX parser content
    xlsx_test = TestXlsxParserConformance()
    doc_xlsx, stream_xlsx = xlsx_test.get_test_document()
    parser_xlsx = XlsxParser()
    ir_xlsx = await parser_xlsx.parse(doc_xlsx, stream_xlsx)
    assert len(ir_xlsx.sections) == 1
    assert ir_xlsx.sections[0].title == "Sheet1"

    # Test CSV parser content
    csv_test = TestCsvParserConformance()
    doc_csv, stream_csv = csv_test.get_test_document()
    parser_csv = CsvParser()
    ir_csv = await parser_csv.parse(doc_csv, stream_csv)
    assert len(ir_csv.sections[0].blocks) == 1
    assert len(ir_csv.sections[0].blocks[0].rows) == 2
